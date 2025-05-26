import os
import logging
from typing import List, Dict, Any
from pathlib import Path
import PyPDF2
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from config.settings import settings

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
        )
    
    def load_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
            return ""
    
    def load_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            return ""
    
    def load_txt(self, file_path: str) -> str:
        """Load text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading TXT {file_path}: {e}")
            return ""
    
    def process_document(self, file_path: str) -> List[Document]:
        """Process a document and return chunks."""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            text = self.load_pdf(file_path)
        elif file_extension == '.docx':
            text = self.load_docx(file_path)
        elif file_extension == '.txt':
            text = self.load_txt(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_extension}")
            return []
        
        if not text.strip():
            logger.warning(f"No text extracted from {file_path}")
            return []
        
        # Split text into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create Document objects with metadata
        documents = []
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk,
                metadata={
                    "source": file_path,
                    "chunk_id": i,
                    "file_name": Path(file_path).name,
                    "file_type": file_extension
                }
            )
            documents.append(doc)
        
        logger.info(f"Processed {file_path}: {len(documents)} chunks created")
        return documents
    
    def process_directory(self, directory_path: str) -> List[Document]:
        """Process all supported documents in a directory."""
        all_documents = []
        supported_extensions = {'.pdf', '.docx', '.txt'}
        
        directory = Path(directory_path)
        if not directory.exists():
            logger.error(f"Directory {directory_path} does not exist")
            return []
        
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                documents = self.process_document(str(file_path))
                all_documents.extend(documents)
        
        logger.info(f"Processed {len(all_documents)} total document chunks")
        return all_documents